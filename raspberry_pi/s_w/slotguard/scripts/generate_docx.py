#!/usr/bin/env python3
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
DOCUMENTS = [
    (
        BASE_DIR / "docs" / "SLOT-GUARD_4inch_LCD_UIUX_기획서_v0.2.md",
        BASE_DIR / "docs" / "SLOT-GUARD_4inch_LCD_UIUX_기획서_v0.2.docx",
    ),
    (
        BASE_DIR / "docs" / "ATmega128A_UART_구현명세_v1.0.md",
        BASE_DIR / "docs" / "ATmega128A_UART_구현명세_v1.0.docx",
    ),
    (
        BASE_DIR
        / "docs"
        / "SLOT-GUARD_Raspberry_Pi_전체_구현명세_v1.0.md",
        BASE_DIR
        / "docs"
        / "SLOT-GUARD_Raspberry_Pi_전체_구현명세_v1.0.docx",
    ),
]

FONT_NAME = "Noto Sans CJK KR"
MONO_FONT_NAME = "D2Coding"
NAVY = "132A44"
BLUE = "0B5B8E"
LIGHT_BLUE = "EAF4FA"
LIGHT_GRAY = "F1F5F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_run_font(run, name=FONT_NAME, size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("- ")
    set_run_font(run, size=9, color="64748B")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])
    end_run = paragraph.add_run(" -")
    set_run_font(end_run, size=9, color="64748B")


def configure_document(
    document,
    header_text="약SLOT-GUARD · UI/UX SPECIFICATION",
):
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    heading_specs = {
        "Title": (26, NAVY),
        "Heading 1": (18, NAVY),
        "Heading 2": (14, BLUE),
        "Heading 3": (11, BLUE),
    }
    for style_name, (size, color) in heading_specs.items():
        style = document.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(header_text)
    set_run_font(header_run, size=8, bold=True, color="64748B")
    set_page_number(section.footer.paragraphs[0])


def add_rich_text(paragraph, text, base_bold=False):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, MONO_FONT_NAME, size=9, color="9A3412")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, bold=base_bold)


def add_code_block(document, lines):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.35)
    paragraph.paragraph_format.right_indent = Cm(0.35)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    properties.append(shading)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, MONO_FONT_NAME, size=8.5, color="0F172A")


def parse_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(document, table_lines):
    rows = [parse_table_row(line) for line in table_lines]
    if len(rows) >= 2 and all(
        re.fullmatch(r":?-{3,}:?", cell.replace(" ", ""))
        for cell in rows[1]
    ):
        rows.pop(1)
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = True

    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        if row_index == 0:
            set_repeat_table_header(row)
        for column_index in range(column_count):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            value = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_rich_text(paragraph, value, base_bold=row_index == 0)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                if row_index == 0:
                    run.font.color.rgb = RGBColor.from_string(WHITE)
            set_cell_shading(cell, NAVY if row_index == 0 else (
                LIGHT_BLUE if row_index % 2 == 0 else WHITE
            ))
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_title(
    document,
    text,
    subtitle_text="PRODUCT · UI/UX SPECIFICATION",
):
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(50)
    paragraph.paragraph_format.space_after = Pt(26)
    add_rich_text(paragraph, text)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(subtitle_text)
    set_run_font(run, size=11, bold=True, color=BLUE)
    document.add_paragraph().add_run()


def markdown_to_docx(source_path, output_path):
    lines = source_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    is_pi_system_spec = "Raspberry_Pi_전체_구현명세" in source_path.name
    header_text = (
        "약SLOT-GUARD · RASPBERRY PI SYSTEM SPECIFICATION"
        if is_pi_system_spec
        else "약SLOT-GUARD · UI/UX SPECIFICATION"
    )
    subtitle_text = (
        "RASPBERRY PI · SYSTEM IMPLEMENTATION SPECIFICATION"
        if is_pi_system_spec
        else "PRODUCT · UI/UX SPECIFICATION"
    )
    configure_document(document, header_text=header_text)

    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index].rstrip()

        if not line:
            index += 1
            continue

        if line.startswith("```"):
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(document, code_lines)
            index += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            if first_heading and level == 1:
                add_title(
                    document,
                    heading_text,
                    subtitle_text=subtitle_text,
                )
                first_heading = False
            else:
                paragraph = document.add_paragraph(
                    style=f"Heading {min(level, 3)}"
                )
                add_rich_text(paragraph, heading_text)
            index += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", line)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, bullet_match.group(1))
            index += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", line)
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            add_rich_text(paragraph, number_match.group(1))
            index += 1
            continue

        paragraph_lines = [line.rstrip("  ")]
        index += 1
        while index < len(lines):
            candidate = lines[index].rstrip()
            if (
                not candidate
                or candidate.startswith("#")
                or candidate.startswith("```")
                or candidate.startswith("|")
                or re.match(r"^[-*]\s+", candidate)
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate.rstrip("  "))
            index += 1

        paragraph = document.add_paragraph()
        add_rich_text(paragraph, " ".join(paragraph_lines))

    properties = document.core_properties
    properties.title = lines[0].lstrip("# ") if lines else source_path.stem
    properties.subject = (
        "약SLOT-GUARD Raspberry Pi 전체 시스템 구현 문서"
        if is_pi_system_spec
        else "약SLOT-GUARD 4인치 LCD 및 UART 구현 문서"
    )
    properties.author = "약SLOT-GUARD Team"
    properties.keywords = "SLOT-GUARD, Raspberry Pi, ATmega128A, UART, LCD"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(output_path)


def main():
    for source_path, output_path in DOCUMENTS:
        markdown_to_docx(source_path, output_path)


if __name__ == "__main__":
    main()
